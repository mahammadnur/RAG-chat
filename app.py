import streamlit as st
from PyPDF2 import PdfReader
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import streamlit as st
from PyPDF2 import PdfReader
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document
import datetime

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document
import datetime
import io
import base64
from pathlib import Path

# Page configuration
st.set_page_config(
    page_title="📄 Smart PDF Chat Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI - Blue/White theme for dark/light mode compatibility
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 50%, #60a5fa 100%);
        color: white;
        border-radius: 15px;
        margin-bottom: 1rem;
        box-shadow: 0 8px 32px rgba(59, 130, 246, 0.3);
    }
    
    .chat-container {
        background: rgba(248, 250, 252, 0.8);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(148, 163, 184, 0.2);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.1);
    }
    
    @media (prefers-color-scheme: dark) {
        .chat-container {
            background: rgba(15, 23, 42, 0.8);
            border: 1px solid rgba(148, 163, 184, 0.2);
            color: white;
        }
    }
    
    .chat-message {
        padding: 1rem 1.2rem;
        border-radius: 18px;
        margin: 1rem 0;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        animation: fadeIn 0.3s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(10px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .user-message {
        background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%);
        color: white;
        margin-left: 15%;
        border-bottom-right-radius: 5px;
    }
    
    .bot-message {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        color: #1e293b;
        margin-right: 15%;
        border: 1px solid rgba(148, 163, 184, 0.3);
        border-bottom-left-radius: 5px;
    }
    
    @media (prefers-color-scheme: dark) {
        .bot-message {
            background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
            color: #f1f5f9;
            border: 1px solid rgba(148, 163, 184, 0.3);
        }
    }
    
    .input-section {
        background: rgba(255, 255, 255, 0.9);
        backdrop-filter: blur(10px);
        border: 2px solid #e2e8f0;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.05);
    }
    
    @media (prefers-color-scheme: dark) {
        .input-section {
            background: rgba(15, 23, 42, 0.9);
            border: 2px solid rgba(148, 163, 184, 0.3);
            color: white;
        }
    }
    
    .stats-container {
        display: flex;
        justify-content: space-around;
        margin: 1rem 0;
        gap: 1rem;
    }
    
    .stat-box {
        background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
        color: white;
        padding: 1rem;
        border-radius: 12px;
        text-align: center;
        min-width: 120px;
        box-shadow: 0 4px 15px rgba(59, 130, 246, 0.3);
        flex: 1;
    }
    
    .no-conversation {
        text-align: center;
        padding: 3rem 2rem;
        color: #64748b;
        background: rgba(248, 250, 252, 0.5);
        border: 2px dashed #cbd5e1;
        border-radius: 15px;
        margin: 1rem 0;
    }
    
    @media (prefers-color-scheme: dark) {
        .no-conversation {
            background: rgba(15, 23, 42, 0.5);
            color: #94a3b8;
            border: 2px dashed #475569;
        }
    }
    
    /* Hide default streamlit elements */
    .stApp > header {
        background-color: transparent;
    }
    
    .stApp {
        background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
    }
    
    @media (prefers-color-scheme: dark) {
        .stApp {
            background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        }
    }
</style>
""", unsafe_allow_html=True)

# Load environment variables
load_dotenv()

# Configure Gemini API
api_key = os.getenv("Gemini_API_KEY")
if not api_key:
    st.error("⚠️ Please set your Gemini_API_KEY in the .env file")
    st.stop()

genai.configure(api_key=api_key)

# Initialize session state
if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = []
if "processed_files" not in st.session_state:
    st.session_state.processed_files = []
if "vector_store_ready" not in st.session_state:
    st.session_state.vector_store_ready = False

def get_text_from_pdf(pdf_files):
    """Extract text from PDF files with better error handling"""
    text = ""
    successful_files = []
    
    if not pdf_files:
        return text, successful_files
    
    for pdf_file in pdf_files:
        try:
            pdf_reader = PdfReader(pdf_file)
            file_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        file_text += page_text + "\n"
                except Exception as e:
                    st.warning(f"⚠️ Could not read page {page_num + 1} of {pdf_file.name}")
                    continue
            
            if file_text.strip():
                text += f"\n--- Content from {pdf_file.name} ---\n" + file_text
                successful_files.append(pdf_file.name)
            else:
                st.warning(f"⚠️ No text found in {pdf_file.name}")
                
        except Exception as e:
            st.error(f"❌ Error reading PDF file {pdf_file.name}: {str(e)}")
            continue
    
    return text.strip(), successful_files

def create_text_chunks(text):
    """Split text into chunks with improved parameters"""
    if not text:
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    return text_splitter.split_text(text)

def create_vector_store(chunks):
    """Create FAISS vector store with error handling"""
    if not chunks:
        raise ValueError("No text chunks to process.")
    
    try:
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=api_key
        )
        vector_store = FAISS.from_texts(chunks, embeddings)
        vector_store.save_local("faiss_index")
        return vector_store
    except Exception as e:
        st.error(f"❌ Error creating vector store: {str(e)}")
        raise

def get_conversation_chain():
    """Create conversation chain with improved prompt"""
    prompt_template = """
    You are a helpful AI assistant analyzing documents. Answer the question based on the provided context.
    
    Guidelines:
    - Provide detailed, accurate answers based on the context
    - If the answer is not in the context, say "I don't have information about that in the uploaded documents."
    - Use a conversational, friendly tone
    - Structure your response clearly
    - Cite specific information when possible
    
    Context: {context}
    
    Question: {question}
    
    Answer:
    """
    
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0.3,
            google_api_key=api_key
        )
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt)
        return chain
    except Exception as e:
        st.error(f"❌ Error creating conversation chain: {str(e)}")
        raise

def process_question(user_question):
    """Process user question and return response"""
    try:
        if not st.session_state.vector_store_ready:
            st.error("📤 Please upload and process PDF files first!")
            return None
        
        if not os.path.exists("faiss_index"):
            st.error("📤 Vector store not found. Please reprocess your files.")
            return None
        
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=api_key
        )
        
        vector_store = FAISS.load_local(
            "faiss_index",
            embeddings,
            allow_dangerous_deserialization=True
        )
        
        docs = vector_store.similarity_search(user_question, k=5)
        
        if not docs:
            return "I couldn't find relevant information in the uploaded documents for your question."
        
        chain = get_conversation_chain()
        response = chain(
            {"input_documents": docs, "question": user_question},
            return_only_outputs=True
        )
        
        return response['output_text']
        
    except Exception as e:
        st.error(f"❌ Error processing question: {str(e)}")
        return None

def create_pdf_download(conversation_history):
    """Create PDF file from conversation history"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, 'PDF Chat Conversation History', 0, 1, 'C')
        pdf.ln(10)
        
        # Add timestamp
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1)
        pdf.ln(10)
        
        # Add conversations
        pdf.set_font("Arial", size=12)
        for i, (question, answer) in enumerate(conversation_history, 1):
            # Question
            pdf.set_font("Arial", 'B', 12)
            pdf.multi_cell(0, 8, f"Q{i}: {question}")
            pdf.ln(3)
            
            # Answer
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 8, f"A{i}: {answer}")
            pdf.ln(8)
        
        # Return PDF as bytes
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        st.error(f"❌ Error creating PDF: {str(e)}")
        return None

def create_doc_download(conversation_history):
    """Create DOCX file from conversation history with Unicode support"""
    try:
        doc = Document()

        # Title
        title = doc.add_heading('PDF Chat Conversation History', 0)
        title.alignment = 1

        # Timestamp
        doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # Conversations
        for i, (question, answer) in enumerate(conversation_history, 1):
            doc.add_heading(f"Question {i}", level=2)
            doc.add_paragraph(question)
            doc.add_heading(f"Answer {i}", level=2)
            doc.add_paragraph(answer)
            doc.add_paragraph("")  # spacing

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    except Exception as e:
        st.error(f"❌ Error creating DOC: {str(e)}")
        return None


# In the main() function, update as follows:
def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🧠 Curiosity AI</h1>
        <p>Chat with Multiple PDFs - Extract Insights from Documents</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([3, 1])

    with col1:
        # 💬 Input section always at top
        st.markdown("""<div class="input-section">""", unsafe_allow_html=True)

        with st.form("question_form", clear_on_submit=True):
            user_question = st.text_input(
                "💬 Ask me anything about your uploaded documents:",
                placeholder="e.g., What are the main objectives in the report?",
                label_visibility="collapsed"
            )
            submitted = st.form_submit_button("🧠 Ask")

            if submitted and user_question:
                with st.spinner("🤔 Thinking..."):
                    response = process_question(user_question)
                    if response:
                        st.session_state.conversation_history.append((user_question, response))
                        st.rerun()  # rerun to show updated chat with input at top

        st.markdown("</div>", unsafe_allow_html=True)

        # 📜 Conversation (latest first)
        if st.session_state.conversation_history:
            st.markdown("""<div class="chat-container">""", unsafe_allow_html=True)

            for question, answer in reversed(st.session_state.conversation_history):
                st.markdown(f"""
                <div class="chat-message user-message">
                    {question}
                </div>
                <div class="chat-message bot-message">
                    {answer}
                </div>
                """, unsafe_allow_html=True)

            st.markdown("</div>", unsafe_allow_html=True)

    # 📊 Stats
    with col2:
        if st.session_state.conversation_history or st.session_state.processed_files:
            st.markdown("""<div class="stats-container">""", unsafe_allow_html=True)

            if st.session_state.conversation_history:
                st.markdown(f"""
                <div class="stat-box">
                    <h3>{len(st.session_state.conversation_history)}</h3>
                    <p>Questions</p>
                </div>
                """, unsafe_allow_html=True)

            if st.session_state.processed_files:
                st.markdown(f"""
                <div class="stat-box">
                    <h3>{len(st.session_state.processed_files)}</h3>
                    <p>Files</p>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("</div>", unsafe_allow_html=True)

    # 📁 Sidebar
    with st.sidebar:
        st.markdown("## 📁 Document Upload")
        pdf_docs = st.file_uploader(
            "Choose PDF files",
            type=["pdf"],
            accept_multiple_files=True,
            help="Upload one or more PDF files (resumes, reports, etc.)"
        )

        if pdf_docs:
            st.markdown("### 📋 Selected Files:")
            for i, pdf in enumerate(pdf_docs, 1):
                file_size = len(pdf.getvalue()) / 1024
                st.write(f"{i}. {pdf.name} ({file_size:.1f} KB)")

        if st.button("🚀 Process Documents", type="primary", use_container_width=True):
            if not pdf_docs:
                st.error("Please upload at least one PDF file!")
            else:
                with st.spinner("Processing documents..."):
                    progress_bar = st.progress(0)
                    progress_bar.progress(25)
                    raw_text, successful_files = get_text_from_pdf(pdf_docs)

                    if not raw_text:
                        st.error("No text could be extracted from the PDF files!")
                        return

                    progress_bar.progress(50)
                    text_chunks = create_text_chunks(raw_text)

                    if not text_chunks:
                        st.error("No text chunks could be created!")
                        return

                    progress_bar.progress(75)
                    try:
                        create_vector_store(text_chunks)
                        st.session_state.vector_store_ready = True
                        st.session_state.processed_files = successful_files
                        progress_bar.progress(100)

                        st.success(f"✅ Successfully processed {len(successful_files)} files!")
                        st.balloons()
                        st.info(f"📊 Extracted {len(raw_text)} characters in {len(text_chunks)} chunks")
                    except Exception as e:
                        st.error(f"Error processing files: {str(e)}")

        if st.session_state.conversation_history:
            st.markdown("---")
            st.markdown("## 💾 Download Options")

            if st.button("📄 Download as PDF", use_container_width=True):
                pdf_data = create_pdf_download(st.session_state.conversation_history)
                if pdf_data:
                    filename = f"conversation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                    st.download_button(
                        label="⬇️ Click to Download PDF",
                        data=pdf_data,
                        file_name=filename,
                        mime="application/pdf",
                        use_container_width=True
                    )

            if st.button("📝 Download as DOC", use_container_width=True):
                doc_data = create_doc_download(st.session_state.conversation_history)
                if doc_data:
                    filename = f"conversation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    st.download_button(
                        label="⬇️ Click to Download DOC",
                        data=doc_data,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

            if st.button("🗑️ Clear Conversation", use_container_width=True):
                st.session_state.conversation_history = []
                st.rerun()

        with st.expander("📖 How to Use"):
            st.markdown("""
            1. **Upload PDFs**  
            2. **Click ‘Process Documents’**  
            3. **Ask questions (input stays on top)**  
            4. **Latest chats appear at the top**  
            5. **Download as PDF or DOC**
            """)
            st.markdown("""
    <hr style="margin-top: 3rem; margin-bottom: 1rem; border: none; height: 1px; background: #ccc;" />
    <div style='text-align: center; color: #6b7280; font-size: 0.9rem;'>
        © 2025 <strong>Curiosity AI</strong> | Built for Curiosity by Noor
    </div>
    """, unsafe_allow_html=True)

# Run the app
if __name__ == "__main__":
    main()
